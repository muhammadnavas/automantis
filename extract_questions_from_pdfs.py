import json
import os
import re
from pathlib import Path

from pypdf import PdfReader
from docx import Document


DATA_DIR = Path("data")
OUTPUT_FILE = DATA_DIR / "questions_from_pdf.json"


def infer_topic(question: str) -> str:
    q = question.lower()
    if "work" in q:
        return "Time and Work"
    if "train" in q or "distance" in q or "speed" in q:
        return "Time and Distance"
    if "interest" in q:
        return "Simple Interest"
    if "profit" in q or "loss" in q or "selling" in q or "cost price" in q:
        return "Profit and Loss"
    if "%" in q or "percent" in q or "percentage" in q:
        return "Percentages"
    return "General Aptitude"


def normalize_option_text(text: str) -> str:
    text = text.encode("utf-8", "ignore").decode("utf-8", "ignore")
    text = re.sub(r"^[A-Da-d][\).:\-]\s*", "", text).strip()
    return re.sub(r"\s+", " ", text)


def clean_text(text: str) -> str:
    return text.encode("utf-8", "ignore").decode("utf-8", "ignore")


def is_noise_question(text: str) -> bool:
    t = text.strip().lower()
    if not t:
        return True
    noise_prefixes = [
        "direction",
        "directions",
        "set of",
        "questions:",
        "all questions carry",
        "section-",
    ]
    return any(t.startswith(prefix) for prefix in noise_prefixes)


def parse_options_text(text: str) -> list[str]:
    text = clean_text(re.sub(r"\s+", " ", text)).strip()
    if not text:
        return []

    # Support both "a) ... b) ..." and cases where first option is unlabeled before b)
    label_matches = list(re.finditer(r"\b([A-Da-d])[).]\s*", text))
    options: list[str] = []

    if not label_matches:
        return []

    first_label = label_matches[0].group(1).lower()
    if first_label != "a":
        leading = text[: label_matches[0].start()].strip(" |;,")
        if leading:
            options.append(normalize_option_text(leading))

    for idx, match in enumerate(label_matches):
        start = match.end()
        end = label_matches[idx + 1].start() if idx + 1 < len(label_matches) else len(text)
        option_text = text[start:end].strip(" |;,")
        option_text = normalize_option_text(option_text)
        if option_text:
            options.append(option_text)

    return options


def extract_answer_map(raw_text: str) -> dict[int, int]:
    answer_map: dict[int, int] = {}
    for raw_line in raw_text.splitlines():
        line = clean_text(raw_line).strip()
        if not line:
            continue

        pairs = re.findall(r"\b(\d{1,3})\s*[-.):]\s*([A-Da-d])\b", line)
        if len(pairs) < 2 and "answer" not in line.lower() and "key" not in line.lower():
            continue

        for num, choice in pairs:
            answer_map[int(num)] = ord(choice.upper()) - ord("A")
    return answer_map


def parse_question_blocks(
    text: str,
    answer_map: dict[int, int] | None = None,
    keep_without_answer: bool = True,
) -> list[dict]:
    answer_map = answer_map or extract_answer_map(text)
    flat = re.sub(r"\s+", " ", text)

    questions = []
    q_pattern = re.compile(
        r"(?P<num>\d{1,3})\.\s*(?P<body>.*?)(?=(?:\b\d{1,3}\.\s)|$)",
        re.DOTALL,
    )

    for match in q_pattern.finditer(flat):
        num = int(match.group("num"))
        body = match.group("body").strip()
        if len(body) < 20:
            continue

        option_start = re.search(r"\b[A-Da-d][)\.:\-]\s*", body)
        if not option_start:
            continue

        question_text = clean_text(body[: option_start.start()]).strip(" :-")
        if not question_text or len(question_text) > 450:
            continue
        if is_noise_question(question_text):
            continue

        options = parse_options_text(body[option_start.start() :])

        if len(options) < 2:
            continue
        if any(len(opt) > 160 for opt in options):
            continue

        correct_option = answer_map.get(num)
        if correct_option is not None and (correct_option < 0 or correct_option >= len(options)):
            correct_option = None
        if correct_option is None and not keep_without_answer:
            continue

        questions.append(
            {
                "question": question_text,
                "options": options,
                "correct_option": correct_option,
            }
        )

    clean = []
    seen = set()
    for item in questions:
        if item["correct_option"] is not None and (
            item["correct_option"] < 0 or item["correct_option"] >= len(item["options"])
        ):
            continue
        key = (item["question"], tuple(item["options"]), item["correct_option"])
        if key in seen:
            continue
        seen.add(key)

        topic = infer_topic(item["question"])
        explanation = "Solve carefully and verify your answer."
        if topic == "Time and Work":
            explanation = (
                "Use LCM method: assume total work as LCM of individual times, then add daily work rates."
            )

        clean.append(
            {
                "topic": topic,
                "question": item["question"],
                "options": item["options"],
                "correct_option": item["correct_option"],
                "explanation": explanation,
                "answer_available": item["correct_option"] is not None,
            }
        )

    return clean


def extract_from_pdf(pdf_path: Path) -> list[dict]:
    reader = PdfReader(str(pdf_path))
    full_text = []
    for page in reader.pages:
        full_text.append(page.extract_text() or "")
    text = "\n".join(full_text)
    return parse_question_blocks(text, keep_without_answer=True)


def parse_docx_table_questions(document: Document, answer_map: dict[int, int]) -> list[dict]:
    questions: list[dict] = []
    seen: set[tuple[str, tuple[str, ...]]] = set()

    for table in document.tables:
        for row in table.rows:
            cells = [clean_text(cell.text).strip() for cell in row.cells if clean_text(cell.text).strip()]
            if len(cells) < 2:
                continue

            num_match = re.match(r"^(\d{1,3})\.?$", cells[0])
            if not num_match:
                continue

            q_num = int(num_match.group(1))
            body = re.sub(r"\s+", " ", cells[1])
            if "|" in body:
                parts = [part.strip() for part in body.split("|") if part.strip()]
                if len(parts) >= 2:
                    question_text = parts[0]
                    options_text = " ".join(parts[1:])
                else:
                    question_text = body
                    options_text = ""
            else:
                option_start = re.search(r"\b[A-Da-d][)\.:\-]\s*", body)
                if not option_start:
                    continue
                question_text = body[: option_start.start()].strip()
                options_text = body[option_start.start() :]

            question_text = clean_text(question_text).strip(" :-")
            if is_noise_question(question_text) or len(question_text) < 8:
                continue

            options = parse_options_text(options_text)
            if len(options) < 2:
                continue

            key = (question_text, tuple(options))
            if key in seen:
                continue
            seen.add(key)

            correct_option = answer_map.get(q_num)
            if correct_option is not None and (correct_option < 0 or correct_option >= len(options)):
                correct_option = None

            topic = infer_topic(question_text)
            explanation = "Solve carefully and verify your answer."
            if topic == "Time and Work":
                explanation = (
                    "Use LCM method: assume total work as LCM of individual times, then add daily work rates."
                )

            questions.append(
                {
                    "topic": topic,
                    "question": question_text,
                    "options": options,
                    "correct_option": correct_option,
                    "explanation": explanation,
                    "answer_available": correct_option is not None,
                }
            )

    return questions


def extract_from_docx(docx_path: Path) -> list[dict]:
    document = Document(str(docx_path))
    lines = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    lines.append(text)

    full_text = "\n".join(lines)
    answer_map = extract_answer_map(full_text)

    table_questions = parse_docx_table_questions(document, answer_map)
    paragraph_questions = parse_question_blocks(
        full_text,
        answer_map=answer_map,
        keep_without_answer=True,
    )

    merged = []
    seen = set()
    for item in table_questions + paragraph_questions:
        key = (item["question"], tuple(item["options"]))
        if key in seen:
            continue
        seen.add(key)
        if "answer_available" not in item:
            item["answer_available"] = item.get("correct_option") is not None
        merged.append(item)

    return merged


def main() -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    pdf_files = sorted(DATA_DIR.glob("*.pdf"))
    docx_files = sorted(DATA_DIR.glob("*.docx"))

    all_questions = []
    for pdf_path in pdf_files:
        all_questions.extend(extract_from_pdf(pdf_path))
    for docx_path in docx_files:
        all_questions.extend(extract_from_docx(docx_path))

    with open(OUTPUT_FILE, "w", encoding="utf-8") as file:
        json.dump(all_questions, file, indent=2, ensure_ascii=False)

    with_answer = sum(1 for q in all_questions if q.get("correct_option") is not None)
    without_answer = len(all_questions) - with_answer

    print(
        "Extracted "
        f"{len(all_questions)} questions from {len(pdf_files)} PDF(s) and {len(docx_files)} DOCX file(s) "
        f"into {OUTPUT_FILE}"
    )
    print(f"Questions with answer key: {with_answer}, without answer key: {without_answer}")


if __name__ == "__main__":
    main()
